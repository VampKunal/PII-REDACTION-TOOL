import json
import docx

def create_test_doc():
    doc = docx.Document()
    doc.add_heading("PII Redaction Test Document", 0)
    
    with open("data/ground_truth.json", "r", encoding="utf-8") as f:
        ground_truth = json.load(f)
        
    doc.add_heading("Entities", 1)
    for entity in ground_truth:
        text = entity.get("text", "")
        # Add the entity in a paragraph, maybe mixed with some other text or standalone
        doc.add_paragraph(f"Here is some text containing {text} for testing purposes.")
        
    doc.save("data/big_test.docx")
    print("Created data/big_test.docx successfully.")

if __name__ == "__main__":
    create_test_doc()
